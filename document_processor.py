import docx
from PyPDF2 import PdfReader
import io
import subprocess
import tempfile
import os

def split_text(text, max_chunk_size=4000):
    """
    Розділяє текст на менші частини для обробки
    """
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 <= max_chunk_size:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word)

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def extract_text_from_pdf(pdf_file):
    """
    Витягує текст з PDF файлу
    """
    reader = PdfReader(pdf_file)
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return '\n'.join(text)

def extract_text_from_doc(doc_file):
    """
    Витягує текст з .doc файлу використовуючи catdoc
    """
    # Створюємо тимчасовий файл для збереження вмісту
    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
        temp_file.write(doc_file.read())
        temp_path = temp_file.name

    try:
        # Використовуємо catdoc для конвертації .doc в текст
        result = subprocess.run(['catdoc', temp_path], 
                              capture_output=True, 
                              text=True)

        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        else:
            error_msg = result.stderr if result.stderr else "Не вдалося витягнути текст з файлу"
            raise ValueError(error_msg)

    except Exception as e:
        raise ValueError(f"Помилка при обробці .doc файлу: {str(e)}")
    finally:
        # Видаляємо тимчасовий файл
        if os.path.exists(temp_path):
            os.unlink(temp_path)

def extract_text_from_docx(docx_file):
    """
    Витягує текст з .docx файлу
    """
    doc = docx.Document(docx_file)
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Також витягуємо текст з таблиць
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)

def extract_text(file):
    """
    Витягує текст з файлу в залежності від його формату
    """
    file_content = file.read()
    file.seek(0)  # Reset file pointer

    # Визначаємо формат файлу за розширенням
    filename = file.name.lower()

    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(io.BytesIO(file_content))
    elif filename.endswith('.doc'):
        return extract_text_from_doc(io.BytesIO(file_content))
    else:
        raise ValueError("Непідтримуваний формат файлу. Підтримуються формати: .pdf, .docx, .doc")