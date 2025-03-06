from docx import Document
import io

def download_results(analysis_results, selected_types=None):
    """
    Форматування результатів аналізу для завантаження
    """
    output = []

    # Додавання заголовку
    output.append("ЗВІТ АНАЛІЗУ ЮРИДИЧНОГО ДОКУМЕНТА")
    output.append("=" * 50)

    if selected_types is None or len(selected_types) == 0:
        # Загальний аналіз за запитом
        if "general" in analysis_results:
            output.append("\nЗАГАЛЬНИЙ АНАЛІЗ")
            output.append("-" * 30)
            output.append(analysis_results["general"])
    else:
        # Аналіз за вибраними типами
        headers = {
            "risks": "АНАЛІЗ РИЗИКІВ",
            "responsibility": "АНАЛІЗ ВІДПОВІДАЛЬНОСТІ",
            "obligations": "АНАЛІЗ ДОГОВІРНИХ ЗОБОВ'ЯЗАНЬ",
            "compliance": "АНАЛІЗ ВІДПОВІДНОСТІ ЗАКОНОДАВСТВУ",
            "financial": "АНАЛІЗ ФІНАНСОВИХ УМОВ"
        }

        for analysis_type in selected_types:
            if analysis_type in analysis_results:
                output.append(f"\n{headers[analysis_type]}")
                output.append("-" * 30)
                output.append(analysis_results[analysis_type])

    return "\n".join(output)

def create_docx_results(analysis_results, selected_types=None):
    """
    Створення DOCX документу з результатами аналізу
    """
    doc = Document()

    # Додаємо заголовок
    doc.add_heading('ЗВІТ АНАЛІЗУ ЮРИДИЧНОГО ДОКУМЕНТА', 0)

    if selected_types is None or len(selected_types) == 0:
        # Загальний аналіз за запитом
        if "general" in analysis_results:
            doc.add_heading('ЗАГАЛЬНИЙ АНАЛІЗ', 1)
            doc.add_paragraph(analysis_results["general"])
    else:
        # Аналіз за вибраними типами
        headers = {
            "risks": "АНАЛІЗ РИЗИКІВ",
            "responsibility": "АНАЛІЗ ВІДПОВІДАЛЬНОСТІ",
            "obligations": "АНАЛІЗ ДОГОВІРНИХ ЗОБОВ'ЯЗАНЬ",
            "compliance": "АНАЛІЗ ВІДПОВІДНОСТІ ЗАКОНОДАВСТВУ",
            "financial": "АНАЛІЗ ФІНАНСОВИХ УМОВ"
        }

        for analysis_type in selected_types:
            if analysis_type in analysis_results:
                doc.add_heading(headers[analysis_type], 1)
                doc.add_paragraph(analysis_results[analysis_type])

    # Зберігаємо документ в байтовий потік
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return docx_stream